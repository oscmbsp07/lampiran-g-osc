import io
import math
import os
import re
import base64
import datetime as dt
import zipfile
from dataclasses import dataclass
from typing import Dict, List, Optional, Tuple, Set
from concurrent.futures import ThreadPoolExecutor, as_completed
import xml.etree.ElementTree as ET

import streamlit as st
from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.enum.section import WD_ORIENTATION, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


# ============================================================
# CONFIG
# ============================================================
st.set_page_config(page_title="Lampiran G Unit OSC", layout="wide")

ALLOWED_SHEETS = {
    "SERENTAK",
    "PKM",
    "TKR-GUNA",
    "TKR",
    "PKM TUKARGUNA",
    "BGN",
    "BGN EVCB",
    "EVCB",
    "EV",
    "TELCO",
    "PS",
    "SB",
    "CT",
    "PL",
    "KTUP",
    "JP",
    "LJUP",
}

# Tapisan agenda TERHAD (ikut arahan terbaru user)
# NOTE: EV termasuk dalam cluster EVCB (kadang kertas maklumat guna "EV" sahaja).
AGENDA_FILTER_SHEETS = {
    "SERENTAK",
    "PKM",
    "TKR-GUNA",        # termasuk variasi "TG" melalui canonical mapping
    "PKM TUKARGUNA",
    "BGN",
    "TELCO",
    "EVCB",
    "BGN EVCB",
    "EV",
}

DAERAH_ORDER = {"SPU": 0, "SPS": 1, "SPT": 2}

KNOWN_CODES = [
    "PKM", "TKR-GUNA", "TKR", "124A", "204D", "PS", "SB", "CT",
    "KTUP", "LJUP", "JP", "PL",
    "BGN", "EVCB", "EV", "TELCO",
]

PB_CODES = {"PKM", "TKR-GUNA", "TKR", "124A", "204D", "PS", "SB", "CT"}
KEJ_CODES = {"KTUP", "LJUP", "JP"}
JL_CODES = {"PL"}

# UT rules kekal (boleh refine kemudian jika perlu)
UT_ALLOWED_SHEETS = {"SERENTAK", "PKM", "BGN", "BGN EVCB", "TKR-GUNA", "PKM TUKARGUNA", "TKR"}
SERENTAK_UT_ALLOWED_INDUK = {"PB", "PKM", "BGN"}


# ============================================================
# UI HELPERS (BACKGROUND + CSS)
# ============================================================
def _inject_bg_and_css(img_path: str) -> bool:
    try:
        with open(img_path, "rb") as f:
            data = f.read()
    except Exception:
        data = None

    b64 = base64.b64encode(data).decode("utf-8") if data else ""
    ext = os.path.splitext(img_path)[1].lower().replace(".", "")
    if ext in {"jpg", "jpeg"}:
        mime = "image/jpeg"
    elif ext == "png":
        mime = "image/png"
    else:
        mime = "image/*"

    bg_css = ""
    if data:
        bg_css = f"""
        .stApp::before {{
            content: "";
            position: fixed;
            inset: 0;
            z-index: -2;
            background-image:
                linear-gradient(rgba(0,0,0,0.48), rgba(0,0,0,0.48)),
                url("data:{mime};base64,{b64}");
            background-size: cover;
            background-position: center center;
            background-repeat: no-repeat;
            background-attachment: fixed;
            transform: translateZ(0);
        }}
        """

    css = f"""
    <style>
      html, body {{ height: 100%; }}
      body {{
        overflow-y: scroll;
        scrollbar-width: none;
        -ms-overflow-style: none;
      }}
      body::-webkit-scrollbar {{
        width: 0px;
        height: 0px;
        background: transparent;
      }}
      header, footer {{
        visibility: hidden;
        height: 0;
      }}
      .stApp {{
        background: transparent !important;
      }}
      {bg_css}
      section.main > div.block-container {{
        max-width: 1200px;
        padding-top: 0.8rem;
        padding-bottom: 0.8rem;
      }}
      .app-title {{
        text-align: center;
        font-weight: 900;
        letter-spacing: 1px;
        margin: 0.9rem 0 0.2rem 0;
        text-transform: uppercase;
        color: white;
        text-shadow: 0px 2px 14px rgba(0,0,0,0.55);
      }}
      .hero-spacer {{ height: 22vh; }}
      div[data-testid="stVerticalBlockBorderWrapper"] {{
        background: rgba(0,0,0,0.44) !important;
        border: 1px solid rgba(255,255,255,0.12) !important;
        border-radius: 18px !important;
        padding: 14px 16px 12px 16px !important;
        box-shadow: 0 10px 30px rgba(0,0,0,0.25);
        backdrop-filter: blur(2px);
      }}
      h1 a, h2 a, h3 a {{ display: none !important; }}
      label {{ font-size: 0.85rem !important; }}
      .stTextInput input {{ height: 2.35rem !important; }}
      div.stButton > button {{
        width: 100%;
        border-radius: 14px;
        font-weight: 800;
        letter-spacing: 0.8px;
      }}
    </style>
    """
    st.markdown(css, unsafe_allow_html=True)
    return bool(data)


def _parse_ddmmyyyy(s: str) -> Optional[dt.date]:
    s = (s or "").strip()
    if not s:
        return None
    try:
        return dt.datetime.strptime(s, "%d/%m/%Y").date()
    except Exception:
        return None


# ============================================================
# UTIL - NORMALISASI & PARSING
# ============================================================
def is_nan(v) -> bool:
    return v is None or (isinstance(v, float) and math.isnan(v)) or (isinstance(v, str) and v.strip().lower() == "nan")


def clean_fail_no(v) -> str:
    """Untuk parsing/dedup dalaman: buang whitespace keras."""
    if is_nan(v):
        return ""
    s = str(v)
    s = re.sub(r"[\s\r\n\t]+", "", s)
    return s.strip()


def clean_str(v) -> str:
    if is_nan(v):
        return ""
    return str(v).strip()


def is_blankish_text(v) -> bool:
    if v is None or is_nan(v):
        return True
    s = str(v).strip()
    if s == "":
        return True
    s2 = s.lower()
    if s2 in {"-", "—", "–", "n/a", "na", "nil", "tiada"}:
        return True
    if re.fullmatch(r"[-–—\s]+", s):
        return True
    return False


def parse_date_from_cell(val) -> Optional[dt.date]:
    if val is None or (isinstance(val, float) and math.isnan(val)):
        return None
    if isinstance(val, dt.datetime):
        return val.date()
    if isinstance(val, dt.date):
        return val
    if isinstance(val, (int, float)) and 20000 < float(val) < 60000:
        base = dt.date(1899, 12, 30)
        return base + dt.timedelta(days=int(val))

    s = str(val).strip()
    if not s or s.lower() == "nan":
        return None

    m = re.search(r"(\d{1,2})[/-](\d{1,2})[/-](\d{4})", s)
    if m:
        d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
        try:
            return dt.date(y, mo, d)
        except Exception:
            return None

    m = re.search(r"(\d{4})[/-](\d{1,2})[/-](\d{1,2})", s)
    if m:
        y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
        try:
            return dt.date(y, mo, d)
        except Exception:
            return None

    return None


def parse_induk_code(val) -> str:
    if val is None or is_nan(val):
        return ""
    s = str(val).strip()
    if not s:
        return ""
    toks = re.findall(r"[A-Z]{2,5}", s.upper())
    return toks[-1] if toks else ""


def in_range(d: Optional[dt.date], start: dt.date, end: dt.date) -> bool:
    return d is not None and start <= d <= end


def normalize_osc_prefix(s: str) -> str:
    """Untuk parsing dalaman."""
    if not s:
        return ""
    s2 = str(s).strip()
    s2 = re.sub(r"[\s\r\n\t]+", "", s2)
    s2 = re.sub(r"^(MBPS|MPSP)", "MBSP", s2, flags=re.IGNORECASE)
    s2 = re.sub(r"^M\.?B\.?S\.?P", "MBSP", s2, flags=re.IGNORECASE)
    s2 = re.sub(r"^M\.?B\.?P\.?S", "MBSP", s2, flags=re.IGNORECASE)
    return s2.upper()


def format_fail_no_display(v) -> str:
    """
    Untuk OUTPUT Lampiran G (FAIL NO):
    - mesti kekal penuh seperti 'No. Rujukan OSC' (termasuk PIN.(TG), (SPEED), dsb)
    - hanya kemaskan spacing yang jelas mengganggu (tanpa buang kandungan).
    """
    if is_blankish_text(v):
        return ""
    s = str(v).strip()

    # normalize prefix MBPS/MPSP -> MBSP (tanpa kacau selebihnya)
    s = re.sub(r"^(MBPS|MPSP)", "MBSP", s, flags=re.IGNORECASE)
    s = re.sub(r"^M\.?B\.?S\.?P", "MBSP", s, flags=re.IGNORECASE)
    s = re.sub(r"^M\.?B\.?P\.?S", "MBSP", s, flags=re.IGNORECASE)

    # normalize whitespace
    s = s.replace("\r", "\n")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s).strip()

    # normalize around separators (tidy sahaja)
    s = re.sub(r"\s*/\s*", "/", s)
    s = re.sub(r"\s*-\s*", "-", s)
    s = re.sub(r"\s*\+\s*", " + ", s)

    # buang double space
    s = re.sub(r" {2,}", " ", s).strip()
    return s


def sheet_norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip()).upper()


def canonical_sheet_name(sheet: str) -> str:
    s = sheet_norm(sheet)
    s = s.replace("_", " ")
    s = re.sub(r"\s+", " ", s).strip()

    if s in {"E V"}:
        return "EV"

    if s in {"TG", "TUKARGUNA", "TUKAR GUNA"}:
        return "TKR-GUNA"
    if s in {"TKR GUNA", "TKR-GUNA"}:
        return "TKR-GUNA"
    if s in {"PKM TUKAR GUNA", "PKM TUKARGUNA"}:
        return "PKM TUKARGUNA"
    if s in {"BGN-EVCB", "BGN EVCB"}:
        return "BGN EVCB"

    return s


def extract_tail_only(fail_no: str) -> str:
    s = normalize_osc_prefix(fail_no)
    m = re.search(r"/(\d{3,5})(?:[-A-Z\(]|$)", s)
    return m.group(1) if m else ""


def extract_series_tail_key(fail_no: str) -> str:
    s = normalize_osc_prefix(fail_no)
    m = re.search(r"^MBSP/\d+/([^/]+)/(\d{3,5})", s)
    if not m:
        return ""
    return f"{m.group(1)}|{m.group(2)}"


def extract_osc_head(fail_no: str) -> str:
    s = normalize_osc_prefix(fail_no)
    m = re.search(r"^(MBSP)/(\d+)/([^/]+)/(\d{3,5})", s)
    if not m:
        return ""
    return f"{m.group(1)}/{m.group(2)}/{m.group(3)}/{m.group(4)}"


def osc_norm(x: str) -> str:
    s = normalize_osc_prefix(str(x or ""))
    s = s.lower()
    s = re.sub(r"[\s\r\n\t]+", "", s)
    s = re.sub(r"[-/\\()\[\]{}+.,:;]", "", s)
    return s


def keputusan_is_empty(v) -> bool:
    if v is None or is_nan(v):
        return True
    s = str(v).strip()
    if s == "" or s.lower() in {"-", "tiada", "nil", "n/a", "na", "—", "–"}:
        return True
    if re.fullmatch(r"[-–—\s]+", s):
        return True
    if parse_date_from_cell(s) is not None:
        return False
    return False


def is_serentak(sheet_name: str, fail_no_display_or_raw: str) -> bool:
    if canonical_sheet_name(sheet_name) == "SERENTAK":
        return True
    return "SERENTAK" in str(fail_no_display_or_raw or "").upper()


def _sheet_implied_codes(sheet_u: str) -> Set[str]:
    s = sheet_u.upper()
    out = set()
    if "PKM" in s:
        out.add("PKM")
    if "TKR-GUNA" in s or s == "TKR GUNA" or s == "TG":
        out.add("TKR-GUNA")
    elif re.fullmatch(r"TKR", s):
        out.add("TKR")
    if "BGN" in s:
        out.add("BGN")
    if "EVCB" in s:
        out.add("EVCB")
    if re.fullmatch(r"EV", s):
        out.add("EV")
    if "TELCO" in s:
        out.add("TELCO")
    if re.fullmatch(r"PS", s):
        out.add("PS")
    if re.fullmatch(r"SB", s):
        out.add("SB")
    if re.fullmatch(r"CT", s):
        out.add("CT")
    if re.fullmatch(r"PL", s):
        out.add("PL")
    if re.fullmatch(r"KTUP", s):
        out.add("KTUP")
    if re.fullmatch(r"JP", s):
        out.add("JP")
    if re.fullmatch(r"LJUP", s):
        out.add("LJUP")
    return out


def extract_codes(fail_no: str, sheet_name: str) -> Set[str]:
    s = normalize_osc_prefix(str(fail_no or ""))
    tokens = re.split(r"[\s\+\-/\\(),]+", s.upper())
    codes: Set[str] = set()
    for t in tokens:
        if t in KNOWN_CODES:
            codes.add(t)

    sn = canonical_sheet_name(sheet_name)
    codes |= _sheet_implied_codes(sn)

    if sn in {"BGN EVCB"}:
        codes.add("BGN")
        codes.add("EVCB")
    return codes


def split_fail_induk(fail_no: str) -> str:
    """
    Dapatkan induk (tanpa suffix jenis permohonan).
    Operasi atas versi normalized (tiada whitespace).
    """
    s = normalize_osc_prefix(str(fail_no or "")).strip()
    if not s:
        return s
    for i in range(len(s) - 1, 0, -1):
        if s[i] == "-":
            suffix = s[i + 1:].upper()
            if any(code in suffix for code in KNOWN_CODES):
                return s[:i]
    return s


def canon_serentak_codes(codes: Set[str]) -> List[str]:
    order = [
        "PKM", "TKR-GUNA", "TKR", "124A", "204D", "PS", "SB", "CT",
        "KTUP", "LJUP", "JP",
        "PL",
        "BGN", "EVCB", "EV", "TELCO",
    ]
    return [c for c in order if c in codes]


def perkara_3lines(d: Optional[dt.date]) -> str:
    dd = d.strftime("%d.%m.%Y") if d else ""
    return f"Penyediaan Kertas\nMesyuarat Tamat Tempoh\n{dd}"


def extract_jenis_from_fail_no_display(fail_no_display: str) -> str:
    """
    RULE UTAMA (ikut arahan user):
    Jenis Permohonan di Lampiran G = ambil bahagian hujung selepas '/NNNN-' (dash selepas nombor tail).
    Contoh:
      MBSP/15/U6-2601/0295-PKM PIN.(TG) + BGN PIN.(TG)
      -> PKM PIN.(TG) + BGN PIN.(TG)
    """
    s = format_fail_no_display(fail_no_display)
    if not s:
        return ""
    m = re.search(r"/\d{3,5}-(.+)$", s)
    if not m:
        return ""
    tail = (m.group(1) or "").strip()
    tail = re.sub(r"^\s*[-–—]\s*", "", tail).strip()
    return tail


# ============================================================
# DISPLAY FORMATTER (Proper Case + keep acronym)
# ============================================================
_ROMAN = {"i","ii","iii","iv","v","vi","vii","viii","ix","x","xi","xii","xiii","xiv","xv","xvi","xvii","xviii","xix","xx"}

def format_pemohon_display(name: str) -> str:
    """
    Format 'PEMAJU/PEMOHON' untuk output Lampiran G (data mentah biasanya ALL CAPS):
    - Default: Proper Case (contoh: PULAU -> Pulau, HOTEL -> Hotel, MAPLE -> Maple)
    - Kekalkan ALL CAPS HANYA untuk akronim/initialism/kod yang benar-benar munasabah:
        * Allowlist (_KEEP_ACRONYMS_COMMON)
        * Token mengandungi digit (cth 4G, A12)
        * Token ALL CAPS panjang <= 3 (cth KB, HHM, PDC) kecuali stopword (DI/KE/OF/...)
        * Token seperti M&E / R&D / A/C (token bergabung dengan & /)
    - Singkatan korporat (Sdn/Bhd/Berhad dll) dipeta ke bentuk standard, bukan ALL CAPS.
    Nota: Fungsi ini fokus pada kemasan visual output Word; tiada rephrase/ubah susunan.
    """
    if is_blankish_text(name):
        return ""

    raw = str(name).replace("\r", "\n")

    # Corporate/suffix (bukan akronim)
    corp_token_map = {
        "SDN": "Sdn",
        "BHD": "Bhd",
        "BERHAD": "Berhad",
        "ENTERPRISE": "Enterprise",
        "ENTERPRISES": "Enterprises",
        "LTD": "Ltd",
        "LIMITED": "Limited",
        "CO": "Co",
        "COMPANY": "Company",
        "INC": "Inc",
        "THE": "The",
        "HOLDINGS": "Holdings",
        "PROPERTIES": "Properties",
        "DEVELOPMENT": "Development",
        "INDUSTRIES": "Industries",
        "CONSORTIUM": "Consortium",
        "MANAGEMENT": "Management",
        "ELECTRONICS": "Electronics",
        "HOTEL": "Hotel",
    }

    # Gelaran (bukan akronim)
    non_acronym_title_map = {
        "TETUAN": "Tetuan",
        "DATO": "Dato",
        "DATO'": "Dato'",
        "DATO’": "Dato’",
        "DATUK": "Datuk",
        "HAJI": "Haji",
        "HAJAH": "Hajah",
        "TUAN": "Tuan",
        "PUAN": "Puan",
        "ENCIK": "Encik",
        "CIK": "Cik",
        "BIN": "bin",
        "BINTI": "binti",
        "BT": "Bt",
        "BTE": "Bte",
        "DR": "Dr",
        "PROF": "Prof",
        "IR": "Ir",
        "TS": "Ts",
        "HJ": "Hj",
        "PN": "Pn",
        "EN": "En",
    }

    # Stopwords (jangan kekal ALL CAPS walaupun 2–3 huruf)
    stopwords_upper = {
        "DI", "KE", "DAN", "ATAU", "DARI", "PADA", "UNTUK", "DALAM", "DENGAN", "ATAS", "BAWAH",
        "OF", "THE", "AND", "OR", "IN", "ON", "AT", "BY", "TO", "FOR", "FROM", "WITH",
    }

    # Regex token: perkataan/alnum termasuk gabungan & atau /
    word_re = re.compile(r"[A-Za-zÀ-ÿ0-9]+(?:[&/][A-Za-zÀ-ÿ0-9]+)*")

    def _is_mixed_case(tok: str) -> bool:
        return any(ch.islower() for ch in tok) and any(ch.isupper() for ch in tok)

    def _format_simple_token(tok: str, is_first_token: bool) -> str:
        if not tok:
            return tok

        if _is_mixed_case(tok):
            return tok  # sengaja mixed-case

        up = tok.upper()
        clean = re.sub(r"[^A-Z0-9]", "", up)

        # Protected: email/url (jarang dalam pemohon, tapi fail-safe)
        if "@" in tok or "://" in tok:
            return tok

        # Nombor semata-mata
        if clean.isdigit():
            return clean

        # Roman numerals
        if clean.lower() in _ROMAN:
            return clean.upper()

        # Corporate mapping
        if clean in corp_token_map:
            return corp_token_map[clean]

        # Title mapping
        if clean in non_acronym_title_map:
            return non_acronym_title_map[clean]

        # Allowlist akronim
        if clean in _KEEP_ACRONYMS_COMMON:
            return clean

        # Token bergabung seperti M&E, R&D, A/C -> kekal ALL CAPS
        if "&" in tok or "/" in tok:
            # format setiap subtoken, separator kekal
            parts = re.split(r"([&/])", tok)
            out_parts = []
            for p in parts:
                if p in {"&", "/"}:
                    out_parts.append(p)
                else:
                    out_parts.append(_format_simple_token(p, is_first_token))
            return "".join(out_parts)

        # Jika asal ALL CAPS:
        if tok == up:
            # stopword -> lower, kecuali token pertama (biar Title Case)
            if clean in stopwords_upper:
                return clean.lower() if not is_first_token else clean.lower().capitalize()

            # Ada digit + huruf -> kekal
            if any(ch.isdigit() for ch in clean) and any(ch.isalpha() for ch in clean):
                return clean

            # Initialism pendek (<=3) -> kekal (cth KB, HHM, PDC)
            if re.fullmatch(r"[A-Z]{1,3}", clean):
                return clean

            # Selain itu, anggap perkataan biasa -> Proper Case
            return tok.lower().capitalize()

        # Default: Proper Case
        return tok.lower().capitalize()

    def _format_line(line: str) -> str:
        s = (line or "")

        # kekalkan spacing/punctuation asal semampu mungkin; cuma buang trailing whitespace
        s = s.strip()

        # Apply per token
        idx = 0
        out = []
        first = True
        for m in word_re.finditer(s):
            out.append(s[idx:m.start()])
            tok = m.group(0)
            out.append(_format_simple_token(tok, first))
            first = False
            idx = m.end()
        out.append(s[idx:])

        # kemaskan multiple space yang ekstrem (tanpa ubah line break)
        t = "".join(out)
        t = re.sub(r"[ 	]{2,}", " ", t).strip()
        return t

    lines = [ln for ln in raw.split("\n")]
    return "\n".join([_format_line(ln) for ln in lines]).strip()

def format_mukim_display(s: str) -> str:
    """
    Kolum MUKIM (data mentah biasanya ALL CAPS):
    - Default: Proper Case
    - Fix khas singkatan yang kerap muncul: MK -> Mk, SEK -> Sek (termasuk variasi M.K / S.E.K / M/K)
    - Kekalkan akronim rasmi dalam allowlist (MBSP/OSC/dll) sebagai ALL CAPS
    """
    if is_blankish_text(s):
        return ""

    lines = str(s).replace("\r", "\n").split("\n")

    force_map = {
        "MUKIM": "Mukim",
        "SEKSYEN": "Seksyen",
        "MK": "Mk",
        "SEK": "Sek",
    }

    mk_pat = re.compile(r"(?i)\bM\s*[\./]?\s*K\b")
    sek_pat = re.compile(r"(?i)\bS\s*[\./]?\s*E\s*[\./]?\s*K\b")

    word_re = re.compile(r"[A-Za-zÀ-ÿ0-9]+(?:[&/][A-Za-zÀ-ÿ0-9]+)*")

    def _fmt(tok: str, is_first: bool) -> str:
        if not tok:
            return tok

        if any(ch.islower() for ch in tok) and any(ch.isupper() for ch in tok):
            return tok

        up = tok.upper()
        clean = re.sub(r"[^A-Z0-9]", "", up)

        if clean.isdigit():
            return clean

        if clean.lower() in _ROMAN:
            return clean.upper()

        if clean in _KEEP_ACRONYMS_COMMON:
            return clean

        if clean in force_map:
            return force_map[clean]

        # stopwords ringkas
        if tok == up and clean in {"DI", "KE", "DAN", "ATAU", "OF", "THE", "AND", "OR", "IN", "ON", "AT", "BY", "TO"}:
            return clean.lower() if not is_first else clean.lower().capitalize()

        # initialism pendek (<=3) kekal
        if tok == up and re.fullmatch(r"[A-Z]{1,3}", clean):
            return clean

        return tok.lower().capitalize()

    def _fmt_line(line: str) -> str:
        if is_blankish_text(line):
            return ""
        t = line.strip()
        # fix MK/SEK walau ada dot/slash/space
        t = mk_pat.sub("Mk", t)
        t = sek_pat.sub("Sek", t)

        idx = 0
        out = []
        first = True
        for m in word_re.finditer(t):
            out.append(t[idx:m.start()])
            out.append(_fmt(m.group(0), first))
            first = False
            idx = m.end()
        out.append(t[idx:])
        t2 = "".join(out)
        t2 = re.sub(r"[ \t]{2,}", " ", t2).strip()
        return t2

    return "\n".join([_fmt_line(ln) for ln in lines if ln.strip()]).strip()

def format_lot_display(s: str) -> str:
    if is_blankish_text(s):
        return ""
    lines = str(s).replace("\r", "\n").split("\n")
    force = {"LOT": "Lot", "PLOT": "Plot", "NO": "No", "NO.": "No.", "PT": "PT"}
    # NOTE: word boundary tidak cover "NO." sebagai satu token bila ada titik; jadi kita handle selepas.
    out_lines = []
    for ln in lines:
        t = _proper_case_line_with_rules(ln, force, _KEEP_ACRONYMS_COMMON)
        # Fix "No." variants
        t = re.sub(r"\bNo\.\b", "No.", t)
        t = re.sub(r"\bNo\b\s*\.?", "No.", t) if re.search(r"\bNo\b\s*\.", t) else t
        out_lines.append(t)
    return "\n".join(out_lines).strip()


# ============================================================
# UT "Belum memberi" mapper — tambah alias (KEJURUTERAAN, PERANCANG BANDAR, dll)
# ============================================================
def tindakan_ut(belum_text: str) -> str:
    if is_blankish_text(belum_text):
        return ""
    raw = str(belum_text).strip()

    # Split list jabatan (support comma/&//)
    parts = [p.strip() for p in re.split(r"[,&/]+", raw) if p.strip()]

    def _norm_token(x: str) -> str:
        s = (x or "").upper().strip()
        s = re.sub(r"^(JABATAN|BAHAGIAN|UNIT|SEKSYEN)\s+", "", s)
        s = re.sub(r"\s+", "", s)
        s = re.sub(r"[^A-Z0-9]", "", s)
        return s

    internal_map = {
        "KEJ": "Pengarah Kejuruteraan",
        "KEJURUTERAAN": "Pengarah Kejuruteraan",

        "PB": "Pengarah Perancang Bandar",
        "PERANCANGBANDAR": "Pengarah Perancang Bandar",

        "BGN": "Pengarah Bangunan",
        "BANGUNAN": "Pengarah Bangunan",

        "COB": "Pengarah COB",
        "PESURUHJAYABANGUNAN": "Pengarah COB",

        "KES": "Pengarah Kesihatan",
        "KESIHATAN": "Pengarah Kesihatan",

        "PEN": "Pengarah Penilaian",
        "PENILAIAN": "Pengarah Penilaian",

        "PBRN": "Pengarah Perbandaran",
        "PERBANDARAN": "Pengarah Perbandaran",

        "LESEN": "Pengarah Pelesenan",
        "PELESENAN": "Pengarah Pelesenan",

        "JL": "Pengarah Landskap",
        "LANDSKAP": "Pengarah Landskap",
    }

    alias_substrings = [
        ("KEJURUTERAAN", "Pengarah Kejuruteraan"),
        ("PERANCANGBANDAR", "Pengarah Perancang Bandar"),
        ("BANGUNAN", "Pengarah Bangunan"),
        ("PESURUHJAYABANGUNAN", "Pengarah COB"),
        ("KESIHATAN", "Pengarah Kesihatan"),
        ("PENILAIAN", "Pengarah Penilaian"),
        ("PERBANDARAN", "Pengarah Perbandaran"),
        ("PELESENAN", "Pengarah Pelesenan"),
        ("LANDSKAP", "Pengarah Landskap"),
    ]

    internal, external = [], []
    for p in parts:
        if is_blankish_text(p):
            continue
        key = _norm_token(p)
        mapped = internal_map.get(key)

        if mapped is None:
            for sub, title in alias_substrings:
                if sub in key:
                    mapped = title
                    break

        if mapped is not None:
            internal.append(mapped)
        else:
            external.append(p.upper())

    def dedup(seq):
        seen, out = set(), []
        for x in seq:
            if x not in seen:
                seen.add(x)
                out.append(x)
        return out

    internal = dedup(internal)
    external = dedup(external)
    return "\n".join(internal + external).strip()


def pemohon_norm(x: str) -> str:
    s = str(x or "").lower().strip()
    s = re.sub(r"\b(tetuan|tuan|puan)\b", "", s)
    s = re.sub(r"\b(sdn\.?\s*bhd\.?|sdn\s*bhd|bhd|berhad|enterprise|enterprises|plc|llp|ltd)\b", "", s)
    s = re.sub(r"[^a-z0-9]+", "", s)
    return s


def lot_tokens(x: str) -> Set[str]:
    toks = re.findall(r"\d{2,6}", str(x or ""))
    return set(toks)


# ============================================================
# AGENDA PARSER (WORD .docx) — tail-focused
# ============================================================
@dataclass
class AgendaBlock:
    is_ptj: bool
    codes: Set[str]
    osc_heads: List[str]
    tails: Set[str]
    series_tail_keys: Set[str]
    pemohon_key: str
    lot_set: Set[str]
    has_osc: bool


@dataclass
class AgendaIndex:
    tails_all: Set[str]
    series_tail_all: Set[str]
    osc_head_norm_all: Set[str]
    blocks: List[AgendaBlock]


HEADER_ANYWHERE_RE = re.compile(r"(?i)\bKERTAS\s+MESYUARAT\s+BIL\.\s*OSC/")
HEADER_CODE_RE = re.compile(r"(?i)OSC/([A-Z]{2,12}(?:-[A-Z]{2,12})?)/")

OSC_HEAD_RE = re.compile(
    r"(?i)\b(MBSP|MBPS|MPSP)\s*/\s*(\d+)\s*/\s*([A-Z0-9\-]+)\s*/\s*(\d{3,5})\b"
)

NO_RUJ_OSC_LINE_RE = re.compile(r"(?i)No\.?\s*Rujukan\s*OSC\s*:?\s*(.+)")
PEMOHON_LINE_RE = re.compile(r"(?i)Pemohon\s*:?\s*(.+)")


def _docx_collect_text(doc: Document) -> str:
    chunks: List[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            chunks.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    chunks.append(t)
    return "\n".join(chunks)


def _extract_images_from_docx_bytes(file_bytes: bytes) -> List[bytes]:
    out = []
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            for name in z.namelist():
                if name.lower().startswith("word/media/") and name.lower().endswith(
                    (".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff")
                ):
                    out.append(z.read(name))
    except Exception:
        return []
    return out


def _try_ocr_images(images: List[bytes]) -> str:
    if not images:
        return ""
    try:
        import pytesseract  # type: ignore
    except Exception:
        return ""

    texts = []
    for b in images:
        try:
            im = Image.open(io.BytesIO(b)).convert("RGB")
            txt = pytesseract.image_to_string(im, lang="eng")
            if txt and txt.strip():
                texts.append(txt)
        except Exception:
            continue
    return "\n".join(texts)


def _split_into_blocks(full_text: str) -> List[str]:
    lines = full_text.splitlines()
    idx = []
    for i, line in enumerate(lines):
        if HEADER_ANYWHERE_RE.search(line):
            idx.append(i)
    if not idx:
        return [full_text] if full_text.strip() else []

    blocks = []
    for j, start in enumerate(idx):
        end = idx[j + 1] if j + 1 < len(idx) else len(lines)
        blk = "\n".join(lines[start:end]).strip()
        if blk:
            blocks.append(blk)
    return blocks


def _parse_block_codes(header_line: str) -> Set[str]:
    codes = set()
    m = HEADER_CODE_RE.search(header_line)
    if not m:
        return codes
    raw = m.group(1).upper().strip()
    if raw in KNOWN_CODES:
        codes.add(raw)
    if raw in {"BGN-EVCB", "BGN EVCB"}:
        codes.add("BGN")
        codes.add("EVCB")
    return codes


def _parse_agenda_block(block_text: str) -> AgendaBlock:
    lines = block_text.splitlines()
    header_line = ""
    for ln in lines[:3]:
        if HEADER_ANYWHERE_RE.search(ln):
            header_line = ln.strip()
            break
    if not header_line:
        header_line = (lines[0].strip() if lines else "")

    is_ptj = bool(re.search(r"(?i)\bOSC/PTJ/", header_line))
    codes = _parse_block_codes(header_line)

    osc_heads: List[str] = []
    tails: Set[str] = set()
    series_tail_keys: Set[str] = set()
    has_osc = False

    for m in OSC_HEAD_RE.finditer(block_text):
        prefix = normalize_osc_prefix(m.group(1))
        yy = m.group(2)
        series = m.group(3).upper()
        tail = m.group(4)
        head = f"{prefix}/{yy}/{series}/{tail}"
        osc_heads.append(head)
        tails.add(tail)
        series_tail_keys.add(f"{series}|{tail}")
        has_osc = True

    for m in NO_RUJ_OSC_LINE_RE.finditer(block_text):
        rhs = (m.group(1) or "").strip()
        if not rhs:
            continue
        rhs2 = normalize_osc_prefix(rhs)
        if rhs2 in {"-", "—", "–"}:
            continue
        mm = OSC_HEAD_RE.search(rhs2)
        if mm:
            prefix = normalize_osc_prefix(mm.group(1))
            yy = mm.group(2)
            series = mm.group(3).upper()
            tail = mm.group(4)
            head = f"{prefix}/{yy}/{series}/{tail}"
            osc_heads.append(head)
            tails.add(tail)
            series_tail_keys.add(f"{series}|{tail}")
            has_osc = True

    seen = set()
    osc_heads2 = []
    for x in osc_heads:
        if x not in seen:
            seen.add(x)
            osc_heads2.append(x)
    osc_heads = osc_heads2

    pem = ""
    m = PEMOHON_LINE_RE.search(block_text)
    if m:
        pem = (m.group(1) or "").strip()
    else:
        m2 = re.search(r"(?i)\bTetuan\b\s*:?\s*(.+)", block_text)
        if m2:
            pem = (m2.group(1) or "").strip()
    pem_key = pemohon_norm(pem)

    lot_candidates = []
    for mm in re.finditer(r"(?i)\b(?:di\s+atas\s+)?lot\b[^.\n\r]{0,160}", block_text):
        lot_candidates.append(mm.group(0))
    for mm in re.finditer(r"(?i)\bPT\s*\d{1,6}\b", block_text):
        lot_candidates.append(mm.group(0))
    lot_s = " ".join(lot_candidates) if lot_candidates else block_text
    lot_set = set(re.findall(r"\d{2,6}", lot_s))

    return AgendaBlock(
        is_ptj=is_ptj,
        codes=codes,
        osc_heads=osc_heads,
        tails=tails,
        series_tail_keys=series_tail_keys,
        pemohon_key=pem_key,
        lot_set=lot_set,
        has_osc=has_osc,
    )


def parse_agenda_docx(file_bytes: bytes, enable_ocr: bool = False) -> AgendaIndex:
    doc = Document(io.BytesIO(file_bytes))
    text_main = _docx_collect_text(doc)

    text_ocr = ""
    if enable_ocr:
        imgs = _extract_images_from_docx_bytes(file_bytes)
        text_ocr = _try_ocr_images(imgs)

    full_text = (text_main + "\n" + (text_ocr or "")).strip()

    blocks_raw = _split_into_blocks(full_text)
    blocks: List[AgendaBlock] = []

    tails_all: Set[str] = set()
    series_tail_all: Set[str] = set()
    osc_head_norm_all: Set[str] = set()

    for blk_text in blocks_raw:
        blk = _parse_agenda_block(blk_text)
        blocks.append(blk)

        if blk.is_ptj:
            continue

        for t in blk.tails:
            tails_all.add(t)
        for k in blk.series_tail_keys:
            series_tail_all.add(k)
        for h in blk.osc_heads:
            osc_head_norm_all.add(osc_norm(h))

    return AgendaIndex(
        tails_all=tails_all,
        series_tail_all=series_tail_all,
        osc_head_norm_all=osc_head_norm_all,
        blocks=blocks,
    )


# ============================================================
# EXCEL READER (ULTRA FAST XML) + ROBUST PER-ROW FALLBACK
# ============================================================
HEADER_HINTS = [
    "No. Rujukan OSC",
    "No. Rujukan",
    "Rujukan OSC",
    "No Fail Permohonan",
    "Pemaju",
    "Pemohon",
    "Daerah",
    "Mukim",
    "Lot",
    "Tempoh Untuk Proses",
    "Tempoh Untuk Diberi",
    "Tarikh Keputusan",
    "Jabatan Induk / Teknikal",
    "Belum memberi",
    "Jenis Permohonan",
]

COL_CANDIDATES = {
    "fail_no": ["norujukanosc", "no rujukan osc", "rujukan osc", "fail no", "failno", "no rujukan", "nofailpermohonan"],
    "pemohon": ["pemajupemohon", "pemaju/pemohon", "pemaju", "pemohon", "tetuan"],
    "mukim": ["mukimseksyen", "mukim/seksyen", "mukim", "seksyen"],
    "lot": ["lot"],
    "jenis_perm": ["jenispermohonan"],
    "km": ["tempohuntukprosesolehjabataninduk", "tempohuntukproses"],
    "ut": ["tempohuntukdiberiulasanolehjabatanteknikal", "tempohuntukdiberiulasan"],
    "belum": ["jabatanindukteknikalygbelummemberikeputusanulasansehinggakini", "belummemberikeputusanulasan", "ygbelummemberikeputusan", "belummemberi"],
    "keputusan": ["tarikhkeputusankuasa", "tarikhkeputusan"],
}

_NS_MAIN = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}"
_NS_REL = "{http://schemas.openxmlformats.org/package/2006/relationships}"


def norm_basic(s: str) -> str:
    s = "" if s is None else str(s)
    s = s.strip().lower()
    s = re.sub(r"[\s\r\n\t]+", " ", s)
    s = re.sub(r"[^a-z0-9]+", "", s)
    return s


def _col_letters_to_index(col_letters: str) -> int:
    col_letters = col_letters.upper()
    n = 0
    for ch in col_letters:
        if "A" <= ch <= "Z":
            n = n * 26 + (ord(ch) - ord("A") + 1)
    return n - 1


_CELL_REF_RE = re.compile(r"^([A-Z]+)(\d+)$")


def _cell_ref_to_col_idx(cell_ref: str) -> Optional[int]:
    if not cell_ref:
        return None
    m = _CELL_REF_RE.match(cell_ref.upper())
    if not m:
        return None
    return _col_letters_to_index(m.group(1))


def _load_shared_strings(z: zipfile.ZipFile) -> List[str]:
    path = "xl/sharedStrings.xml"
    if path not in z.namelist():
        return []
    data = z.read(path)
    root = ET.fromstring(data)
    out: List[str] = []
    for si in root.findall(f".//{_NS_MAIN}si"):
        texts = []
        for t in si.findall(f".//{_NS_MAIN}t"):
            if t.text:
                texts.append(t.text)
        out.append("".join(texts))
    return out


def _workbook_sheet_paths(z: zipfile.ZipFile) -> List[Tuple[str, str]]:
    wb_xml = z.read("xl/workbook.xml")
    wb_root = ET.fromstring(wb_xml)

    rels_xml = z.read("xl/_rels/workbook.xml.rels")
    rels_root = ET.fromstring(rels_xml)

    rid_to_target: Dict[str, str] = {}
    for rel in rels_root.findall(f".//{_NS_REL}Relationship"):
        rid = rel.attrib.get("Id", "")
        target = rel.attrib.get("Target", "")
        if rid and target:
            if not target.startswith("xl/"):
                target = "xl/" + target.lstrip("/")
            rid_to_target[rid] = target

    out: List[Tuple[str, str]] = []
    sheets_el = wb_root.find(f".//{_NS_MAIN}sheets")
    if sheets_el is None:
        return out

    for sh in sheets_el.findall(f"{_NS_MAIN}sheet"):
        name = sh.attrib.get("name", "")
        rid = sh.attrib.get(f"{{http://schemas.openxmlformats.org/officeDocument/2006/relationships}}id", "")
        target = rid_to_target.get(rid, "")
        if name and target:
            out.append((name, target))
    return out


def _cell_value_from_c_el(c_el: ET.Element, shared_strings: List[str]) -> Optional[object]:
    t = c_el.attrib.get("t", "")
    v_el = c_el.find(f"{_NS_MAIN}v")

    if t == "inlineStr":
        is_el = c_el.find(f"{_NS_MAIN}is")
        if is_el is None:
            return None
        texts = []
        for t_el in is_el.findall(f".//{_NS_MAIN}t"):
            if t_el.text:
                texts.append(t_el.text)
        txt = "".join(texts).strip()
        return txt if txt != "" else None

    if v_el is None or v_el.text is None:
        return None

    raw = v_el.text
    if t == "s":
        try:
            idx = int(raw)
            if 0 <= idx < len(shared_strings):
                return shared_strings[idx]
            return None
        except Exception:
            return None

    if t == "b":
        return True if raw == "1" else False

    s = raw.strip()
    if s == "":
        return None
    if re.fullmatch(r"-?\d+", s):
        try:
            return int(s)
        except Exception:
            return s
    if re.fullmatch(r"-?\d+\.\d+", s):
        try:
            return float(s)
        except Exception:
            return s
    return s


def _iter_sheet_rows_cells(
    z: zipfile.ZipFile,
    sheet_path: str,
    shared_strings: List[str],
    max_rows_to_scan: Optional[int] = None
):
    if sheet_path not in z.namelist():
        return
    with z.open(sheet_path) as f:
        context = ET.iterparse(f, events=("end",))
        yielded = 0
        for _, elem in context:
            if elem.tag == f"{_NS_MAIN}row":
                r_attr = elem.attrib.get("r", "")
                try:
                    rnum = int(r_attr) if r_attr else None
                except Exception:
                    rnum = None

                cells: Dict[int, object] = {}
                for c in elem.findall(f"{_NS_MAIN}c"):
                    ref = c.attrib.get("r", "")
                    col_idx = _cell_ref_to_col_idx(ref)
                    if col_idx is None:
                        continue
                    val = _cell_value_from_c_el(c, shared_strings)
                    if val is None:
                        continue
                    cells[col_idx] = val

                if rnum is not None and cells:
                    yield rnum, cells
                    yielded += 1
                    if max_rows_to_scan is not None and yielded >= max_rows_to_scan:
                        break

                elem.clear()


def _row_cells_to_list(cells: Dict[int, object]) -> List[object]:
    if not cells:
        return []
    mx = max(cells.keys())
    out = [""] * (mx + 1)
    for k, v in cells.items():
        out[k] = v
    return out


def _header_score(joined_lower: str) -> int:
    score = 0
    for h in HEADER_HINTS:
        if h.lower() in joined_lower:
            score += 1
    return score


def _find_header_row_ultra(rows_iter) -> Tuple[Optional[int], Optional[List[object]]]:
    best_r = None
    best_score = 0
    best_vals: Optional[List[object]] = None

    for rnum, cells in rows_iter:
        vals = _row_cells_to_list(cells)
        joined = " | ".join([str(x).strip() for x in vals if str(x).strip()]).lower()
        if not joined:
            continue
        score = _header_score(joined)
        if score > best_score:
            best_score = score
            best_r = rnum
            best_vals = vals

    if best_r is None or best_score == 0:
        return None, None
    return best_r, best_vals


def _detect_columns_candidates(header_vals: List[object]) -> Dict[str, List[int]]:
    norm_cols = [norm_basic(x) for x in header_vals]
    cand: Dict[str, List[int]] = {k: [] for k in COL_CANDIDATES.keys()}

    for key, needles in COL_CANDIDATES.items():
        for idx, ncol in enumerate(norm_cols):
            for needle in needles:
                if needle and needle in ncol:
                    cand[key].append(idx)
                    break
    return cand


# --- ROBUST FALLBACK per-row (avoid "Pemaju/Lot kosong" bila duplicate header) ---
def _is_nonempty(v) -> bool:
    if v is None:
        return False
    if isinstance(v, str):
        s = v.strip()
        if s == "" or s.lower() == "nan":
            return False
        if s.upper() in {"#N/A", "#VALUE!", "#REF!", "#DIV/0!", "#NAME?"}:
            return False
        if s.lower() in {"-", "—", "–", "n/a", "na", "nil", "tiada"}:
            return False
        if re.fullmatch(r"[-–—\s]+", s):
            return False
    return True


def _is_code_like(v) -> bool:
    if not _is_nonempty(v):
        return False
    s = str(v).upper()
    return bool(re.search(r"\b(PKM|TKR|TKR[-\s]?GUNA|124A|204D|PS|SB|CT|KTUP|LJUP|JP|PL|BGN|EVCB|EV|TELCO)\b", s))


def _rank_columns(cand_idxs: List[int], sample_rows: List[Dict[int, object]], prefer_code: bool = False) -> List[int]:
    if not cand_idxs:
        return []
    scored = []
    for idx in cand_idxs:
        score = 0
        for cells in sample_rows:
            v = cells.get(idx)
            if prefer_code:
                if _is_code_like(v):
                    score += 3
                elif _is_nonempty(v):
                    score += 1
            else:
                if _is_nonempty(v):
                    score += 1
        scored.append((score, idx))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [idx for _, idx in scored] if scored else cand_idxs


def _pick_from_cols(cells: Dict[int, object], col_list: List[int]) -> Optional[object]:
    for idx in col_list:
        v = cells.get(idx)
        if _is_nonempty(v):
            return v
    return None


def read_kertas_excel_ultra(excel_bytes: bytes, daerah_label: str) -> List[dict]:
    out: List[dict] = []
    allowed_upper = {s.upper() for s in ALLOWED_SHEETS}

    with zipfile.ZipFile(io.BytesIO(excel_bytes)) as z:
        shared = _load_shared_strings(z)
        sheet_paths = _workbook_sheet_paths(z)

        for sheet_name, sheet_path in sheet_paths:
            sheet_clean = canonical_sheet_name(sheet_name)
            if sheet_clean.upper() not in allowed_upper:
                continue

            scan_iter = _iter_sheet_rows_cells(z, sheet_path, shared, max_rows_to_scan=220)
            hdr_rnum, hdr_vals = _find_header_row_ultra(scan_iter)
            if hdr_rnum is None or hdr_vals is None:
                continue

            cand = _detect_columns_candidates(hdr_vals)

            sample_rows: List[Dict[int, object]] = []
            for rnum, cells in _iter_sheet_rows_cells(z, sheet_path, shared, max_rows_to_scan=None):
                if rnum <= hdr_rnum:
                    continue
                sample_rows.append(cells)
                if len(sample_rows) >= 160:
                    break

            fail_cols = _rank_columns(cand.get("fail_no", []), sample_rows)
            pem_cols = _rank_columns(cand.get("pemohon", []), sample_rows)
            mukim_cols = _rank_columns(cand.get("mukim", []), sample_rows)
            lot_cols = _rank_columns(cand.get("lot", []), sample_rows)
            jenis_cols = _rank_columns(cand.get("jenis_perm", []), sample_rows, prefer_code=True)
            km_cols = _rank_columns(cand.get("km", []), sample_rows)
            ut_cols = _rank_columns(cand.get("ut", []), sample_rows)
            belum_cols = _rank_columns(cand.get("belum", []), sample_rows)
            keputusan_cols = _rank_columns(cand.get("keputusan", []), sample_rows)

            if not fail_cols or not pem_cols:
                continue

            for rnum, cells in _iter_sheet_rows_cells(z, sheet_path, shared, max_rows_to_scan=None):
                if rnum <= hdr_rnum:
                    continue

                fail = _pick_from_cols(cells, fail_cols)
                pem = _pick_from_cols(cells, pem_cols)

                # OUTPUT mesti guna full No. Rujukan OSC (bukan yang dibersihkan)
                fail_disp = format_fail_no_display(clean_str(fail))
                fail_raw_for_parse = normalize_osc_prefix(clean_fail_no(fail))

                pem_str = clean_str(pem)

                if (is_nan(fail) or fail_disp == "") and (is_nan(pem) or pem_str == ""):
                    continue

                mukim_val = _pick_from_cols(cells, mukim_cols) if mukim_cols else None
                lot_val = _pick_from_cols(cells, lot_cols) if lot_cols else None
                jenis_val = _pick_from_cols(cells, jenis_cols) if jenis_cols else None

                km_raw = _pick_from_cols(cells, km_cols) if km_cols else None
                ut_raw = _pick_from_cols(cells, ut_cols) if ut_cols else None
                belum_val = _pick_from_cols(cells, belum_cols) if belum_cols else None
                keputusan_val = _pick_from_cols(cells, keputusan_cols) if keputusan_cols else None

                rec = {
                    "daerah": daerah_label,
                    "sheet": sheet_clean,

                    # untuk output
                    "fail_no_disp": fail_disp,

                    # untuk parsing/dedup internal
                    "fail_no_raw": fail_raw_for_parse,

                    "pemohon": pem_str,
                    "mukim": clean_str(mukim_val) if mukim_val is not None else "",
                    "lot": clean_str(lot_val) if lot_val is not None else "",
                    "jenis_row": clean_str(jenis_val) if jenis_val is not None else "",
                    "km_date": parse_date_from_cell(km_raw) if km_raw is not None else None,
                    "ut_date": parse_date_from_cell(ut_raw) if ut_raw is not None else None,
                    "belum": clean_str(belum_val) if belum_val is not None else "",
                    "keputusan": clean_str(keputusan_val) if keputusan_val is not None else "",
                    "induk_code": parse_induk_code(km_raw),
                }
                out.append(rec)

    return out


@st.cache_data(show_spinner=False)
def cached_read_kertas_excel_ultra(excel_bytes: bytes, daerah_label: str) -> List[dict]:
    return read_kertas_excel_ultra(excel_bytes, daerah_label)


# ============================================================
# BUILD CATEGORIES
# ============================================================
def parse_primary_code(jenis_row: str, sheet_u: str) -> str:
    s = (jenis_row or "").upper().strip()
    if s:
        s = s.replace("TKR GUNA", "TKR-GUNA")
        m = re.search(r"\b(TKR-GUNA|PKM|TKR|124A|204D|PS|SB|CT|KTUP|LJUP|JP|PL|BGN|EVCB|EV|TELCO)\b", s)
        if m:
            return m.group(1)

    su = canonical_sheet_name(sheet_u).upper()
    if su in {"PKM", "TKR", "TKR-GUNA", "KTUP", "JP", "LJUP", "PL", "PS", "SB", "CT", "EVCB", "EV", "TELCO", "BGN"}:
        return su
    if su in {"PKM TUKARGUNA"}:
        return "TKR-GUNA"
    return ""


def enrich_rows(rows: List[dict]) -> List[dict]:
    out = []
    for r in rows:
        rr = dict(r)
        rr["sheet_u"] = canonical_sheet_name(r["sheet"])
        rr["codes"] = extract_codes(r["fail_no_raw"], rr["sheet_u"])
        rr["primary_code"] = parse_primary_code(r.get("jenis_row", ""), rr["sheet_u"])

        # serentak detect guna sheet atau fail_no_disp (lebih tepat ikut arahan)
        rr["serentak"] = is_serentak(rr["sheet_u"], r.get("fail_no_disp", "") or r.get("fail_no_raw", ""))

        rr["fail_induk"] = split_fail_induk(r["fail_no_raw"])

        rr["tail"] = extract_tail_only(r["fail_no_raw"])
        rr["series_tail"] = extract_series_tail_key(r["fail_no_raw"])
        rr["osc_head_norm"] = osc_norm(extract_osc_head(r["fail_no_raw"]))

        rr["pemohon_key"] = pemohon_norm(r.get("pemohon", ""))
        rr["lot_set"] = lot_tokens(r.get("lot", ""))
        out.append(rr)
    return out


def sheet_is_ut_allowed(sheet_u: str) -> bool:
    s = canonical_sheet_name(sheet_u)
    if s in UT_ALLOWED_SHEETS:
        return True
    if "GUNA" in s and ("TKR" in s or "TUKAR" in s or s == "TG"):
        return True
    return False


def _agenda_fallback_match(row: dict, agenda: "AgendaIndex") -> bool:
    if row["sheet_u"] not in AGENDA_FILTER_SHEETS:
        return False
    if not row.get("pemohon_key"):
        return False
    if not row.get("lot_set"):
        return False

    row_codes = set(row.get("codes") or set())

    for blk in agenda.blocks:
        if blk.is_ptj:
            continue
        if not blk.pemohon_key or not blk.lot_set:
            continue

        if blk.codes and not (row_codes & blk.codes):
            continue

        if row["pemohon_key"] != blk.pemohon_key:
            continue

        inter = row["lot_set"] & blk.lot_set
        if not inter:
            continue

        if min(len(row["lot_set"]), len(blk.lot_set)) >= 2 and len(inter) < 2:
            continue

        return True

    return False


def build_categories(
    rows: List[dict],
    agenda: Optional["AgendaIndex"],
    km_start: dt.date,
    km_end: dt.date,
    ut_start: dt.date,
    ut_end: dt.date,
    ut_enabled: bool,
    agenda_enabled: bool,
) -> Tuple[List[dict], List[dict], List[dict], List[dict], List[dict]]:

    rows = [r for r in rows if keputusan_is_empty(r.get("keputusan"))]

    if agenda_enabled and agenda:
        def _keep(r: dict) -> bool:
            if r["sheet_u"] not in AGENDA_FILTER_SHEETS:
                return True

            if r.get("tail") and r["tail"] in agenda.tails_all:
                return False
            if r.get("series_tail") and r["series_tail"] in agenda.series_tail_all:
                return False
            if r.get("osc_head_norm") and r["osc_head_norm"] in agenda.osc_head_norm_all:
                return False
            if _agenda_fallback_match(r, agenda):
                return False
            return True

        rows = [r for r in rows if _keep(r)]

    by_induk: Dict[str, List[dict]] = {}
    for r in rows:
        by_induk.setdefault(r["fail_induk"], []).append(r)

    def nama_simplify(x: str) -> str:
        return pemohon_norm(x)

    def make_rec(cat: int, tindakan: str, base_r: dict, jenis: str, fail_no: str, perkara: str, extra_key: str) -> dict:
        return {
            "cat": cat,
            "tindakan": tindakan,
            "jenis": jenis,
            "fail_no": fail_no,
            "pemohon": base_r["pemohon"],  # display formatting dibuat masa output Word
            "daerah": base_r["daerah"],
            "mukim": base_r["mukim"],
            "lot": base_r["lot"],
            "perkara": perkara,
            "dedup_key": f"{cat}|{tindakan}|{osc_norm(fail_no)}|{nama_simplify(base_r['pemohon'])}|{extra_key}",
        }

    UT_PRIMARY_ALLOWED = {"PKM", "TKR", "TKR-GUNA", "BGN", "EVCB", "EV", "TELCO"}
    UT_PRIMARY_DISALLOWED = {"KTUP", "LJUP", "JP", "PL", "PS", "SB", "CT", "204D", "124A"}

    cat1, cat2, cat3, cat4, cat5 = [], [], [], [], []

    for induk, grp in by_induk.items():
        is_ser = any(g["serentak"] for g in grp)

        union_codes: Set[str] = set()
        km_dates = [g["km_date"] for g in grp if g.get("km_date")]
        for g in grp:
            union_codes |= set(g["codes"])
        km_date = min(km_dates) if km_dates else None

        # pilih FAIL NO display terbaik (paling panjang biasanya paling lengkap, termasuk (SPEED), PIN.(TG), dsb)
        fail_no_disp_best = ""
        cand_disp = [g.get("fail_no_disp", "") for g in grp if g.get("fail_no_disp")]
        if cand_disp:
            fail_no_disp_best = max(cand_disp, key=lambda x: len(str(x)))
        else:
            fail_no_disp_best = induk

        # JENIS PERMOHONAN ikut rule user: ambil tail selepas /NNNN-
        jenis_best = extract_jenis_from_fail_no_display(fail_no_disp_best)
        if is_ser:
            if jenis_best:
                if "(SERENTAK)" not in jenis_best.upper():
                    jenis_best = f"{jenis_best} (Serentak)"
            else:
                jenis_best = "(Serentak)"

        # KATEGORI 1 — KM
        if is_ser and in_range(km_date, km_start, km_end):
            if union_codes & (PB_CODES - {"PS", "SB", "CT"}):
                cat1.append(make_rec(1, "Pengarah Perancang Bandar", grp[0], jenis_best, fail_no_disp_best, perkara_3lines(km_date), "SER-PB"))
            if union_codes & {"BGN", "EVCB", "EV", "TELCO"}:
                cat1.append(make_rec(1, "Pengarah Bangunan", grp[0], jenis_best, fail_no_disp_best, perkara_3lines(km_date), "SER-BGN"))

        if not is_ser:
            for g in grp:
                if not in_range(g.get("km_date"), km_start, km_end):
                    continue

                fail_no_disp = g.get("fail_no_disp", "") or g["fail_no_raw"]
                jenis = extract_jenis_from_fail_no_display(fail_no_disp) or g["sheet_u"]

                if g["codes"] & {"PKM", "TKR", "TKR-GUNA"}:
                    cat1.append(make_rec(1, "Pengarah Perancang Bandar", g, jenis, fail_no_disp, perkara_3lines(g.get("km_date")), "NS-PB"))
                if g["codes"] & {"BGN", "EVCB", "EV", "TELCO"}:
                    cat1.append(make_rec(1, "Pengarah Bangunan", g, jenis, fail_no_disp, perkara_3lines(g.get("km_date")), "NS-BGN"))

        # KATEGORI 2 — UT (row-level filter by primary_code)
        if ut_enabled:
            for g in grp:
                if not sheet_is_ut_allowed(g["sheet_u"]):
                    continue
                if not in_range(g.get("ut_date"), ut_start, ut_end):
                    continue
                if is_blankish_text(g.get("belum")):
                    continue

                pc = (g.get("primary_code") or "").upper().strip()
                if pc in UT_PRIMARY_DISALLOWED:
                    continue
                if pc and pc not in UT_PRIMARY_ALLOWED:
                    continue

                if g["sheet_u"] == "SERENTAK":
                    if pc:
                        if pc not in UT_PRIMARY_ALLOWED:
                            continue
                    else:
                        if (g.get("induk_code") or "") and (g.get("induk_code") not in SERENTAK_UT_ALLOWED_INDUK):
                            continue
                        if not (set(g.get("codes") or set()) & UT_PRIMARY_ALLOWED):
                            continue

                tindakan = tindakan_ut(g.get("belum", ""))
                if is_blankish_text(tindakan):
                    continue

                fail_no_disp = g.get("fail_no_disp", "") or g["fail_no_raw"]
                jenis = extract_jenis_from_fail_no_display(fail_no_disp) or (jenis_best if is_ser else g["sheet_u"])
                if is_ser and "(SERENTAK)" not in jenis.upper():
                    jenis = f"{jenis} (Serentak)" if jenis else "(Serentak)"

                perkara = f"Ulasan teknikal belum dikemukakan. Tamat Tempoh {g['ut_date'].strftime('%d.%m.%Y')}."
                extra_key = f"{g['sheet_u']}|{pc}|{g['ut_date'].isoformat()}|{(g.get('belum') or '').strip()}"
                cat2.append(make_rec(2, tindakan, g, jenis, fail_no_disp, perkara, extra_key))

        # KATEGORI 3/4/5 — KM
        if is_ser and in_range(km_date, km_start, km_end):
            if union_codes & KEJ_CODES:
                cat3.append(make_rec(3, "Pengarah Kejuruteraan", grp[0], jenis_best, fail_no_disp_best, perkara_3lines(km_date), "SER-KEJ"))
            if union_codes & JL_CODES:
                cat4.append(make_rec(4, "Pengarah Landskap", grp[0], jenis_best, fail_no_disp_best, perkara_3lines(km_date), "SER-JL"))
            if union_codes & {"124A", "204D"}:
                cat5.append(make_rec(5, "Pengarah Perancang Bandar", grp[0], jenis_best, fail_no_disp_best, perkara_3lines(km_date), "SER-124A204D"))

        if not is_ser:
            for g in grp:
                if in_range(g.get("km_date"), km_start, km_end) and (g["sheet_u"] in {"KTUP", "JP", "LJUP"}):
                    fail_no_disp = g.get("fail_no_disp", "") or g["fail_no_raw"]
                    jenis = extract_jenis_from_fail_no_display(fail_no_disp) or g["sheet_u"]
                    cat3.append(make_rec(3, "Pengarah Kejuruteraan", g, jenis, fail_no_disp, perkara_3lines(g.get("km_date")), f"NS-{g['sheet_u']}"))

                if in_range(g.get("km_date"), km_start, km_end) and (g["sheet_u"] == "PL"):
                    fail_no_disp = g.get("fail_no_disp", "") or g["fail_no_raw"]
                    jenis = extract_jenis_from_fail_no_display(fail_no_disp) or g["sheet_u"]
                    cat4.append(make_rec(4, "Pengarah Landskap", g, jenis, fail_no_disp, perkara_3lines(g.get("km_date")), "NS-PL"))

                if in_range(g.get("km_date"), km_start, km_end) and (g["sheet_u"] in {"PS", "SB", "CT"}):
                    fail_no_disp = g.get("fail_no_disp", "") or g["fail_no_raw"]
                    jenis = extract_jenis_from_fail_no_display(fail_no_disp) or g["sheet_u"]
                    cat5.append(make_rec(5, "Pengarah Perancang Bandar", g, jenis, fail_no_disp, perkara_3lines(g.get("km_date")), f"NS-{g['sheet_u']}"))

    def dedup_list(lst: List[dict]) -> List[dict]:
        seen, out2 = set(), []
        for r in lst:
            if r["dedup_key"] in seen:
                continue
            seen.add(r["dedup_key"])
            out2.append(r)
        return out2

    cat1, cat2, cat3, cat4, cat5 = map(dedup_list, [cat1, cat2, cat3, cat4, cat5])

    cat1.sort(key=lambda r: (0 if r["tindakan"].startswith("Pengarah Perancang") else 1, DAERAH_ORDER.get(r["daerah"], 9), r["fail_no"]))
    cat2.sort(key=lambda r: (DAERAH_ORDER.get(r["daerah"], 9), r["fail_no"], r["tindakan"]))
    cat3.sort(key=lambda r: (DAERAH_ORDER.get(r["daerah"], 9), r["fail_no"]))
    cat4.sort(key=lambda r: (DAERAH_ORDER.get(r["daerah"], 9), r["fail_no"]))
    cat5.sort(key=lambda r: (DAERAH_ORDER.get(r["daerah"], 9), r["fail_no"]))

    for lst in [cat1, cat2, cat3, cat4, cat5]:
        for i, r in enumerate(lst, start=1):
            r["bil"] = i

    return cat1, cat2, cat3, cat4, cat5


# ============================================================
# WORD FORMATTER
# ============================================================
COL_WIDTHS_IN = [0.50, 1.75, 1.55, 1.55, 1.45, 0.75, 0.65, 0.70, 1.99]
HEADERS = ["BIL", "TINDAKAN", "JENIS\nPERMOHONAN", "FAIL NO", "PEMAJU/PEMOHON", "DAERAH", "MUKIM", "LOT", "PERKARA"]


def _find_font_path(prefer_bold: bool = True) -> Optional[str]:
    candidates = [
        "/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman_Bold.ttf",
        "/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
    ]
    for p in candidates:
        if os.path.exists(p) and (not prefer_bold or "Bold" in os.path.basename(p)):
            return p
    for p in candidates:
        if os.path.exists(p):
            return p
    return None


def make_g_logo_png(diameter_px: int = 140, outline_px: int = 4, font_pt: int = 34) -> bytes:
    scale = 4
    D = diameter_px * scale
    img = Image.new("RGBA", (D, D), (255, 255, 255, 0))
    dr = ImageDraw.Draw(img)

    pad = (outline_px + 4) * scale
    dr.ellipse((pad, pad, D - pad, D - pad), outline=(0, 0, 0, 255), width=outline_px * scale)

    font_path = _find_font_path(prefer_bold=True)
    size_px = int(font_pt * 96 / 72) * scale
    if font_path:
        try:
            font = ImageFont.truetype(font_path, size_px)
        except Exception:
            font = ImageFont.load_default()
    else:
        font = ImageFont.load_default()

    bbox = dr.textbbox((0, 0), "G", font=font)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    x = (D - tw) / 2 - bbox[0]
    y = (D - th) / 2 - bbox[1] - int(2 * scale)
    dr.text((x, y), "G", font=font, fill=(0, 0, 0, 255))

    img_small = img.resize((diameter_px, diameter_px), resample=Image.LANCZOS)
    buf = io.BytesIO()
    img_small.save(buf, format="PNG")
    return buf.getvalue()


def set_section_landscape(sec):
    sec.orientation = WD_ORIENTATION.LANDSCAPE
    sec.page_width = Inches(11.69)
    sec.page_height = Inches(8.27)
    sec.left_margin = Inches(0.4)
    sec.right_margin = Inches(0.4)
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)


def clear_header(hdr):
    for p in list(hdr.paragraphs):
        p._element.getparent().remove(p._element)


def add_logo_first_page(sec, logo_png_bytes: bytes):
    sec.different_first_page_header_footer = True
    sec.header.is_linked_to_previous = False
    sec.first_page_header.is_linked_to_previous = False
    sec.footer.is_linked_to_previous = False
    sec.first_page_footer.is_linked_to_previous = False

    clear_header(sec.header)
    clear_header(sec.first_page_header)

    hdr = sec.first_page_header
    p = hdr.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.add_picture(io.BytesIO(logo_png_bytes), width=Inches(0.70))


def set_paragraph_font(p, font_name: str, size_pt: float, bold: bool = False, align=None):
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1
    for r in p.runs:
        r.font.name = font_name
        r._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        r.font.size = Pt(size_pt)
        r.font.bold = bold


def add_blank(doc: Document):
    p = doc.add_paragraph("")
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1


def add_title_line_main(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1

    before = "PERMOHONAN YANG "
    mid = "TELAH DAN AKAN SAMPAI"
    after = " TEMPOH & ULASAN JABATAN TEKNIKAL BELUM DIKEMUKAKAN"

    r1 = p.add_run(before)
    r2 = p.add_run(mid)
    r3 = p.add_run(after)

    for r in [r1, r2, r3]:
        r.font.name = "Trebuchet MS"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Trebuchet MS")
        r.font.bold = True
        r.font.size = Pt(12)

    r2.font.size = Pt(13.5)


def add_center_bold(doc: Document, text: str, font: str = "Trebuchet MS", size: float = 12):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(p, font, size, bold=True)


def add_km_line(doc: Document, km_start: dt.date, km_end: dt.date):
    p = doc.add_paragraph(f"KERTAS MESYUARAT (TEMPOH {km_start.strftime('%d/%m/%Y')} HINGGA {km_end.strftime('%d/%m/%Y')})")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_font(p, "Trebuchet MS", 12, bold=True)
    add_blank(doc)


def add_ut_line(doc: Document, ut_start: dt.date, ut_end: dt.date):
    p = doc.add_paragraph(f"ULASAN TEKNIKAL (TEMPOH {ut_start.strftime('%d/%m/%Y')} HINGGA {ut_end.strftime('%d/%m/%Y')})")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_font(p, "Arial", 12, bold=True)
    add_blank(doc)


def set_cell_vcenter(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = tcPr.find(qn("w:vAlign"))
    if vAlign is None:
        vAlign = OxmlElement("w:vAlign")
        tcPr.append(vAlign)
    vAlign.set(qn("w:val"), "center")


def set_row_as_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = trPr.find(qn("w:tblHeader"))
    if tblHeader is None:
        tblHeader = OxmlElement("w:tblHeader")
        trPr.append(tblHeader)
    tblHeader.set(qn("w:val"), "true")


def set_table_borders(tbl):
    tbl_pr = tbl._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    def _edge(tag):
        el = borders.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")

    for t in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        _edge(t)


def format_table(tbl):
    tbl.autofit = False
    set_table_borders(tbl)

    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(COL_WIDTHS_IN[i])
    for i, col in enumerate(tbl.columns):
        col.width = Inches(COL_WIDTHS_IN[i])

    hdr_row = tbl.rows[0]
    set_row_as_header(hdr_row)

    for i, cell in enumerate(hdr_row.cells):
        cell.text = HEADERS[i]
        set_cell_vcenter(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1
        for run in p.runs:
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
            run.font.size = Pt(9)
            run.font.bold = True

    for r in tbl.rows[1:]:
        for c in r.cells:
            set_cell_vcenter(c)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pf = p.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing = 1
                for run in p.runs:
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
                    run.font.size = Pt(9)
                    run.font.bold = False


def fill_table(tbl, recs: List[dict]):
    note_fields = ["bil", "tindakan", "jenis", "fail_no", "pemohon", "daerah", "mukim", "lot", "perkara"]
    for rec in recs:
        row = tbl.add_row()

        vals = []
        for k in note_fields:
            if k == "pemohon":
                vals.append(format_pemohon_display(str(rec.get(k, ""))))
            elif k == "mukim":
                vals.append(format_mukim_display(str(rec.get(k, ""))))
            elif k == "lot":
                vals.append(format_lot_display(str(rec.get(k, ""))))
            else:
                vals.append(str(rec.get(k, "")))

        for i, val in enumerate(vals):
            cell = row.cells[i]
            cell.text = ""
            set_cell_vcenter(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = 1
            run = p.add_run(str(val))
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
            run.font.size = Pt(9)
            run.font.bold = False


def build_word_doc(
    meeting_info: str,
    km_start: dt.date,
    km_end: dt.date,
    ut_start: dt.date,
    ut_end: dt.date,
    cat1: List[dict],
    cat2: List[dict],
    cat3: List[dict],
    cat4: List[dict],
    cat5: List[dict],
    ut_enabled: bool,
) -> bytes:
    logo_png = make_g_logo_png()

    doc = Document()
    set_section_landscape(doc.sections[0])

    def add_category_section(cat_num: int, recs: List[dict]):
        if cat_num == 1:
            sec = doc.sections[0]
        else:
            sec = doc.add_section(WD_SECTION.NEW_PAGE)
            set_section_landscape(sec)

        add_logo_first_page(sec, logo_png)

        if cat_num == 2:
            add_ut_line(doc, ut_start, ut_end)
        else:
            add_title_line_main(doc)
            if cat_num in {3, 4, 5}:
                bagi = {
                    3: "BAGI PELAN KEJURUTERAAN",
                    4: "BAGI PELAN LANDSKAP",
                    5: "BAGI PELAN PS / SB / CT",
                }[cat_num]
                add_center_bold(doc, bagi)
            add_center_bold(doc, meeting_info.strip())
            add_blank(doc)
            add_km_line(doc, km_start, km_end)

        tbl = doc.add_table(rows=1, cols=9)
        format_table(tbl)
        fill_table(tbl, recs)

    add_category_section(1, cat1)
    if ut_enabled:
        add_category_section(2, cat2)
    add_category_section(3, cat3)
    add_category_section(4, cat4)
    add_category_section(5, cat5)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============================================================
# STREAMLIT UI
# ============================================================
bg_ok = _inject_bg_and_css("assets/bg.jpg")
if not bg_ok:
    st.warning("Background tidak dijumpai. Pastikan fail ada di folder assets/ (contoh: assets/bg.jpg).")

st.markdown("<h1 class='app-title'>LAMPIRAN G UNIT OSC</h1>", unsafe_allow_html=True)
st.markdown("<div class='hero-spacer'></div>", unsafe_allow_html=True)

if "running" not in st.session_state:
    st.session_state.running = False

left_col, right_col = st.columns([1.15, 0.85], gap="large")

with left_col:
    with st.container(border=True):
        st.markdown("### Maklumat Mesyuarat")
        st.markdown("**Maklumat mesyuarat**")
        meeting_info = st.text_input("", value="", key="meeting_info", label_visibility="collapsed")

        st.markdown("### Tempoh Kertas Mesyuarat (KM)")
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("**KM Mula (dd/mm/yyyy)**")
            km_mula_str = st.text_input("", value="", key="km_mula", label_visibility="collapsed")
        with c2:
            st.markdown("**KM Akhir (dd/mm/yyyy)**")
            km_akhir_str = st.text_input("", value="", key="km_akhir", label_visibility="collapsed")

        ut_enabled = st.checkbox("Aktifkan Ulasan Teknikal (UT)", value=True)

        ut_mula_str, ut_akhir_str = "", ""
        if ut_enabled:
            st.markdown("### Tempoh Ulasan Teknikal (UT)")
            u1, u2 = st.columns(2)
            with u1:
                st.markdown("**UT Mula (dd/mm/yyyy)**")
                ut_mula_str = st.text_input("", value="", key="ut_mula", label_visibility="collapsed")
            with u2:
                st.markdown("**UT Akhir (dd/mm/yyyy)**")
                ut_akhir_str = st.text_input("", value="", key="ut_akhir", label_visibility="collapsed")

with right_col:
    with st.container(border=True):
        st.markdown("### Muat Naik Fail")

        st.markdown("**Agenda JK OSC (.docx)**")
        agenda_file = st.file_uploader("", type=["docx"], key="agenda_docx", label_visibility="collapsed")

        enable_agenda_ocr = st.checkbox(
            "Agenda scan (OCR) jika perlu",
            value=False,
            help="Jika agenda dalam bentuk gambar/scan. Jika server tiada OCR library, sistem akan teruskan tanpa OCR (tak crash).",
        )

        proceed_without_agenda = st.checkbox(
            "Teruskan tanpa Agenda",
            value=False,
            help="Tick jika agenda belum diterima. Jika tick, sistem jana tanpa tapisan agenda.",
        )

        st.markdown("**Kertas Maklumat (Excel) — SPU/SPS/SPT (boleh upload 1 atau 2 fail setiap daerah)**")
        st.caption("Nota: Hujung/awal tahun boleh jadi 2 fail (tahun lama + tahun baru). Sistem akan gabungkan.")

        st.markdown("**SPU (maks 2 fail)**")
        spu_files = st.file_uploader("", type=["xlsx", "xlsm"], key="spu_multi", label_visibility="collapsed", accept_multiple_files=True)

        st.markdown("**SPS (maks 2 fail)**")
        sps_files = st.file_uploader("", type=["xlsx", "xlsm"], key="sps_multi", label_visibility="collapsed", accept_multiple_files=True)

        st.markdown("**SPT (maks 2 fail)**")
        spt_files = st.file_uploader("", type=["xlsx", "xlsm"], key="spt_multi", label_visibility="collapsed", accept_multiple_files=True)

mid = st.columns([1, 0.55, 1])[1]
with mid:
    gen = st.button("JANA LAMPIRAN G", type="primary", disabled=st.session_state.running)


# ============================================================
# ACTION (GENERATE)
# ============================================================
if gen:
    st.session_state.running = True

    try:
        with st.spinner("Sedang jana Lampiran G..."):
            km_start = _parse_ddmmyyyy(km_mula_str)
            km_end = _parse_ddmmyyyy(km_akhir_str)

            if ut_enabled:
                ut_start = _parse_ddmmyyyy(ut_mula_str)
                ut_end = _parse_ddmmyyyy(ut_akhir_str)
            else:
                ut_start = None
                ut_end = None

            agenda_enabled = True
            if proceed_without_agenda:
                agenda_enabled = False
            else:
                if not agenda_file:
                    st.error("Sila upload Agenda JK OSC (.docx) atau tick 'Teruskan tanpa Agenda'.")
                    st.stop()

            spu_files = spu_files or []
            sps_files = sps_files or []
            spt_files = spt_files or []

            if len(spu_files) == 0 or len(sps_files) == 0 or len(spt_files) == 0:
                st.error("Sila upload sekurang-kurangnya 1 fail untuk setiap SPU, SPS dan SPT.")
                st.stop()

            if len(spu_files) > 2 or len(sps_files) > 2 or len(spt_files) > 2:
                st.error("Maksimum 2 fail dibenarkan bagi setiap daerah (SPU/SPS/SPT).")
                st.stop()

            if km_start is None or km_end is None:
                st.error("Sila isi tarikh KM Mula dan KM Akhir dalam format dd/mm/yyyy.")
                st.stop()
            if km_start > km_end:
                st.error("KM Mula tidak boleh lebih besar daripada KM Akhir.")
                st.stop()

            if ut_enabled:
                if ut_start is None or ut_end is None:
                    st.error("Sila isi tarikh UT Mula dan UT Akhir dalam format dd/mm/yyyy.")
                    st.stop()
                if ut_start > ut_end:
                    st.error("UT Mula tidak boleh lebih besar daripada UT Akhir.")
                    st.stop()

            agenda_index = None
            if agenda_enabled:
                agenda_bytes = agenda_file.getvalue()
                agenda_index = parse_agenda_docx(agenda_bytes, enable_ocr=enable_agenda_ocr)
                if enable_agenda_ocr:
                    try:
                        import pytesseract  # type: ignore
                    except Exception:
                        st.warning("OCR tidak tersedia pada server ini. Sistem teruskan baca agenda tanpa OCR (text sahaja).")

            # --- ULTRA FAST PARALLEL READ (SPU/SPS/SPT) ---
            rows: List[dict] = []

            def _read_one_bytes(b: bytes, daerah: str) -> List[dict]:
                return cached_read_kertas_excel_ultra(b, daerah)

            tasks = []
            with ThreadPoolExecutor(max_workers=6) as ex:
                for f in spu_files:
                    tasks.append(ex.submit(_read_one_bytes, f.getvalue(), "SPU"))
                for f in sps_files:
                    tasks.append(ex.submit(_read_one_bytes, f.getvalue(), "SPS"))
                for f in spt_files:
                    tasks.append(ex.submit(_read_one_bytes, f.getvalue(), "SPT"))

                for fut in as_completed(tasks):
                    rows += fut.result()

            rows = enrich_rows(rows)

            cat1, cat2, cat3, cat4, cat5 = build_categories(
                rows=rows,
                agenda=agenda_index,
                km_start=km_start,
                km_end=km_end,
                ut_start=ut_start if ut_enabled else km_start,
                ut_end=ut_end if ut_enabled else km_end,
                ut_enabled=ut_enabled,
                agenda_enabled=agenda_enabled,
            )

            doc_bytes = build_word_doc(
                meeting_info=meeting_info.strip() if meeting_info.strip() else "JK OSC",
                km_start=km_start,
                km_end=km_end,
                ut_start=ut_start if ut_enabled else km_start,
                ut_end=ut_end if ut_enabled else km_end,
                cat1=cat1,
                cat2=cat2,
                cat3=cat3,
                cat4=cat4,
                cat5=cat5,
                ut_enabled=ut_enabled,
            )

            st.success("Lampiran G berjaya dijana.")
            st.download_button(
                "Muat turun Lampiran G (Word)",
                data=doc_bytes,
                file_name="Lampiran_G.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

            with st.expander("Ringkasan (untuk semakan cepat)"):
                st.write({
                    "Kategori 1": len(cat1),
                    "Kategori 2": len(cat2) if ut_enabled else 0,
                    "Kategori 3": len(cat3),
                    "Kategori 4": len(cat4),
                    "Kategori 5": len(cat5),
                    "Agenda digunakan?": "YA" if agenda_enabled else "TIDAK (Teruskan tanpa Agenda)",
                    "OCR agenda aktif?": "YA" if (agenda_enabled and enable_agenda_ocr) else "TIDAK",
                    "Sheet ditapis agenda": sorted(list(AGENDA_FILTER_SHEETS)),
                    "Rule tapisan agenda": "TAIL-BASED (No Unik Hujung) — PTJ dikecualikan",
                    "Reader Excel": "ULTRA XML + ROBUST FALLBACK (handle duplicate header lama+baru)",
                    "FAIL NO output": "Copy-paste penuh dari No. Rujukan OSC",
                    "Jenis Permohonan output": "Ambil suffix selepas /NNNN- (+ tambah (Serentak) jika serentak)",
                    "Format Mukim/Lot": "Proper Case (Mukim, Lot)",
                })

    except Exception as e:
        st.error("Proses jana Lampiran G gagal. Ini punca ralat (auto keluar):")
        st.exception(e)

    finally:
        st.session_state.running = False
